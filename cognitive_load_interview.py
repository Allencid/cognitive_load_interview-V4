import streamlit as st
import re
import io
from docx import Document

st.set_page_config(page_title="時序陳述分析系統", layout="wide")
st.title("🧭 時序陳述分段與補問工具")

# 初始化階段狀態
if "stage" not in st.session_state:
    st.session_state.stage = "input"
    st.session_state.statement = ""  # 確保陳述內容初始化
    st.session_state.split_method = ""  # 分段方式初始化
    st.session_state.max_segs = 3  # 預設分段數

# 分段函式
def split_text_by_time(text):
    time_cues = r"(?<!\d)(\d{1,2}點(?:\d{1,2}分)?|早上|中午|下午|傍晚|晚上|清晨|凌晨|當時|後來|接著|那時候|之後|突然|隔天|前一天|一天|某天|同時)(?!\d)"
    split_points = [m.start() for m in re.finditer(time_cues, text)]
    segments = []

    if not split_points:
        return [text]  # 無時序詞，整段輸出

    split_points.append(len(text))
    for i in range(len(split_points)-1):
        segment = text[split_points[i]:split_points[i+1]].strip()
        if segment:
            segments.append(segment)

    return segments

if st.session_state.stage == "input":
    st.markdown("""
    請你把件事情從頭到尾仔細地跟我說一遍，不管重要或不重要的細節都跟我說，
    你可以自己決定從什麼時候開始說，也可以自己決定說到什麼時候結束：
    """)
    statement = st.text_area("✏️ 請輸入完整陳述內容：", height=300, value=st.session_state.statement)

    # 讓用戶選擇分段方式
    st.subheader("📐 選擇分段方式")
    split_method = st.radio("選擇分段方式：", ("時間", "語意"))
    st.session_state.split_method = split_method

    # 讓用戶選擇分段數量
    max_segs = st.slider("選擇分段數量：", min_value=1, max_value=10, value=st.session_state.max_segs)
    st.session_state.max_segs = max_segs

    if st.button("📌 確認並開始分段") and statement:
        st.session_state.statement = statement
        st.session_state.stage = "segmentation"
        st.rerun()

elif st.session_state.stage == "segmentation":
    statement = st.session_state.statement
    responses = []

    # 根據分段方式選擇分段數
    max_segs = st.session_state.max_segs

    # 根據選擇的分段方式來處理分段
    if st.session_state.split_method == "時間":
        segments = split_text_by_time(statement)
    else:
        # 預設為語意分段，這裡可以加入語意分析的分段邏輯
        segments = [statement[i:i+round(len(statement)/max_segs)] for i in range(0, len(statement), round(len(statement)/max_segs))]

    edited_segments = []
    st.subheader("📗 分段與可編輯補問區")

    for i, seg in enumerate(segments, 1):
        st.markdown(f"### 📍 段落 {i}")
        edited = st.text_area(f"📝 你可以修改這段內容", value=seg, key=f"edit_seg_{i}")
        edited_segments.append(edited)

        with st.expander(f"🗨️ 補問：關於段落 {i}，你可以講得更仔細一點嗎？"):
            answer = st.text_area(f"你想補充的內容（段落{i}）", key=f"answer_{i}")
            responses.append((f"段落{i}", edited, answer))

            if st.checkbox(f"➕ 我要新增其他問題（段落{i}）", key=f"add_q_{i}"):
                user_q = st.text_input(f"請輸入你的自訂問題：", key=f"user_q_{i}")
                user_a = st.text_area(f"你的回答：", key=f"user_a_{i}")
                responses.append((f"自訂問題 - 段落{i}", user_q, user_a))

    # 補問 1
    st.subheader("🔍 延伸補問區")
    with st.expander("補問 1️⃣：這些事情發生之前的任何事情，你可以跟我說得更仔細一點嗎？"):
        q1 = "在這些事情發生之前，有沒有什麼你覺得相關的事？例如事情的起因、當時的背景、前一天發生什麼..."
        a1 = st.text_area("回答：", key="before_event")
        responses.append(("補問1 - 該事件前的事", q1, a1))

    # 補問 2
    with st.expander("補問 2️⃣：事情發生之後的任何事情，你可以跟我說得更仔細一點嗎？"):
        q2 = "在這些事情發生之後，你做了什麼？或別人做了什麼？例如你有沒有講出去、處理後果、跟誰討論..."
        a2 = st.text_area("回答：", key="after_event")
        responses.append(("補問2 - 該事件後的事", q2, a2))

    # 補問 3
    with st.expander("補問 3️⃣：為了讓所有調查這個案件的人可以更加相信你說的是實話，你還可以想到什麼事情可以跟我說的嗎？任何事情都可以？"):
        q3 = "為了讓所有調查這個案件的人可以更加相信你說的是實話，你還可以想到什麼事情可以跟我說的嗎？任何事情都可以？"
        a3 = st.text_area("回答：", key="credibility")
        responses.append(("補問3 - 鼓勵誠實補充", q3, a3))

    if st.button("📄 顯示所有回應彙整"):
        st.session_state.responses = responses
        st.session_state.stage = "summary"
        st.rerun()

    # 加入返回分段選擇的功能
    if st.button("🔄 回到分段選擇"):
        st.session_state.stage = "input"
        st.rerun()

elif st.session_state.stage == "summary":
    st.subheader("📘 使用者所有補充回應")
    responses = st.session_state.responses
    statement = st.session_state.statement

    txt_output = "【完整陳述內容】\n" + statement + "\n\n"
    doc = Document()
    doc.add_heading("使用者補充回應彙整", level=1)
    doc.add_heading("完整陳述內容", level=2)
    doc.add_paragraph(statement)

    for title, context, ans in responses:
        display_text = f"【{title}】\n> {context}\n✏ 回答：{ans if ans else '（無）'}\n\n"
        st.text(display_text)
        txt_output += display_text

        doc.add_heading(title, level=2)
        doc.add_paragraph(context)
        doc.add_paragraph(f"回答：{ans if ans else '（無）'}")

    txt_bytes = txt_output.encode("utf-8")
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)

    st.download_button("📄 下載為 TXT", data=txt_bytes, file_name="陳述補問彙整.txt", mime="text/plain")
    st.download_button("📘 下載為 DOCX", data=doc_buffer, file_name="陳述補問彙整.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
